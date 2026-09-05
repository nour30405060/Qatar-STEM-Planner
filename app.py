import os
from datetime import datetime
from pathlib import Path

import streamlit as st
from pypdf import PdfReader
from openai import OpenAI
from docx import Document

APP_TITLE = "Q-STEM Planner AI"
BASE_DIR = Path(__file__).parent
KB_DIR = BASE_DIR / "knowledge_base"
RUBRIC_PATH = BASE_DIR / "rubrics" / "stem_lesson_plan_rubric.txt"
OUTPUT_DIR = BASE_DIR / "outputs"
OUTPUT_DIR.mkdir(exist_ok=True)

st.set_page_config(page_title=APP_TITLE, page_icon="🧠", layout="wide")


def get_api_key() -> str | None:
    # Recommended: keep the key in Streamlit secrets or environment variables.
    try:
        key = st.secrets.get("OPENAI_API_KEY", None)
        if key:
            return key
    except Exception:
        pass
    return os.getenv("OPENAI_API_KEY")


def extract_pdf_text(file) -> str:
    try:
        reader = PdfReader(file)
        pages = []
        for i, page in enumerate(reader.pages):
            text = page.extract_text() or ""
            pages.append(f"\n--- Page {i+1} ---\n{text}")
        return "\n".join(pages).strip()
    except Exception as e:
        return f"[PDF extraction error: {e}]"


def load_text_from_path(path: Path) -> str:
    if path.suffix.lower() == ".txt":
        return path.read_text(encoding="utf-8", errors="ignore")
    if path.suffix.lower() == ".pdf":
        try:
            with path.open("rb") as f:
                return extract_pdf_text(f)
        except Exception as e:
            return f"[Could not read {path.name}: {e}]"
    return ""


def load_knowledge_base(max_chars: int = 35000) -> str:
    chunks = []
    for path in sorted(KB_DIR.glob("*")):
        if path.is_file() and path.suffix.lower() in [".pdf", ".txt"]:
            text = load_text_from_path(path)
            if text.strip():
                chunks.append(f"\n=== SOURCE: {path.name} ===\n{text}")
    combined = "\n".join(chunks)
    return combined[:max_chars]


def load_rubric() -> str:
    if RUBRIC_PATH.exists():
        return RUBRIC_PATH.read_text(encoding="utf-8", errors="ignore")
    return "Rubric file not found."


def call_ai(system_prompt: str, user_prompt: str, model: str = "gpt-4o-mini") -> str:
    api_key = get_api_key()
    if not api_key:
        return (
            "⚠️ No API key found. Add OPENAI_API_KEY as an environment variable "
            "or in Streamlit secrets to enable AI generation."
        )
    client = OpenAI(api_key=api_key)
    response = client.chat.completions.create(
        model=model,
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt},
        ],
        temperature=0.35,
    )
    return response.choices[0].message.content


def create_word_report(plan_text: str, evaluation_text: str, metadata: dict) -> Path:
    doc = Document()
    doc.add_heading("Q-STEM Planner AI Report", level=1)
    doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    doc.add_heading("Lesson Metadata", level=2)
    for k, v in metadata.items():
        doc.add_paragraph(f"{k}: {v}")
    doc.add_heading("Generated / Adapted STEM Lesson Plan", level=2)
    for para in plan_text.split("\n"):
        doc.add_paragraph(para)
    doc.add_heading("AI-Assisted Quality Evaluation", level=2)
    for para in evaluation_text.split("\n"):
        doc.add_paragraph(para)
    output_path = OUTPUT_DIR / f"q_stem_plan_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    doc.save(output_path)
    return output_path


st.title("🧠 Q-STEM Planner AI")
st.caption("A research prototype for curriculum-grounded, standards-based STEM lesson co-planning and quality evaluation in Qatar.")

with st.sidebar:
    st.header("Research Prototype Controls")
    model = st.selectbox("AI model", ["gpt-4o-mini", "gpt-4o"], index=0)
    st.info("Keep your API key private. Teachers should use the app interface only.")
    st.markdown("### Knowledge Base")
    kb_files = [p.name for p in KB_DIR.glob("*") if p.is_file()]
    if kb_files:
        st.write(kb_files)
    else:
        st.write("No knowledge-base files found yet.")

st.header("1. Upload Lesson PDF")
uploaded_pdf = st.file_uploader("Upload textbook lesson pages as PDF", type=["pdf"])

lesson_text = ""
if uploaded_pdf:
    with st.spinner("Reading PDF..."):
        lesson_text = extract_pdf_text(uploaded_pdf)
    with st.expander("Preview extracted lesson text"):
        st.text_area("Extracted text", lesson_text[:12000], height=250)

st.header("2. Teacher Input")
col1, col2 = st.columns(2)
with col1:
    grade = st.text_input("Grade level", "Grade 8")
    lesson_title = st.text_input("Lesson title / topic", "")
    duration = st.text_input("Lesson duration", "45 minutes")
    class_size = st.text_input("Number of students", "")
with col2:
    resources = st.text_area("Available resources", "Textbook, worksheets, simple classroom materials, projector")
    student_level = st.selectbox("General student level", ["Mixed ability", "Struggling", "Average", "Advanced"])
    local_context = st.text_input("Preferred local context in Qatar", "Sustainability / environment / daily life in Qatar")
    plan_type = st.selectbox("Plan type", ["STEM lesson", "Mini STEM project", "Engineering design challenge"])

st.header("3. Generate Curriculum-Grounded STEM Plan")
if "generated_plan" not in st.session_state:
    st.session_state.generated_plan = ""
if "evaluation" not in st.session_state:
    st.session_state.evaluation = ""

if st.button("Generate STEM Lesson Plan", type="primary"):
    if not lesson_text:
        st.warning("Please upload lesson PDF pages first.")
    else:
        kb = load_knowledge_base()
        rubric = load_rubric()
        system_prompt = """
You are an expert STEM lesson-planning assistant for science teachers in Qatar.
You are not a generic chatbot. You must operate as a standards-based, curriculum-grounded AI co-planning tool.
Use the uploaded textbook lesson content as the primary source. Use the embedded knowledge base and rubric as supporting standards.
Do not invent unrelated lesson content. If you add STEM connections, explicitly connect them to concepts in the uploaded lesson.
Keep the teacher as the final decision-maker. Produce a draft that can be creatively adapted by the teacher.
"""
        user_prompt = f"""
LESSON METADATA:
Grade: {grade}
Lesson title: {lesson_title}
Duration: {duration}
Class size: {class_size}
Student level: {student_level}
Available resources: {resources}
Preferred Qatar/local context: {local_context}
Plan type: {plan_type}

UPLOADED LESSON CONTENT:
{lesson_text[:25000]}

EMBEDDED OFFICIAL / STANDARDS KNOWLEDGE BASE:
{kb}

QUALITY RUBRIC:
{rubric}

TASK:
Create a high-quality STEM lesson plan with the following sections:
1. Brief analysis of the uploaded lesson content
2. Curriculum-grounded lesson focus
3. STEM real-world problem/challenge
4. Learning outcomes
5. STEM integration map: Science, Technology, Engineering, Mathematics
6. Engineering Design Process steps
7. Lesson sequence with timing
8. Teacher role and student role
9. Questions/prompts for students
10. Assessment plan: formative and summative
11. Student product/performance task
12. Differentiation for struggling and advanced students
13. Resources and safety notes
14. Teacher Creative Adaptation options
15. Notes showing how the plan is grounded in the uploaded PDF and standards knowledge base
"""
        with st.spinner("Generating plan..."):
            st.session_state.generated_plan = call_ai(system_prompt, user_prompt, model=model)

st.text_area("Generated STEM Plan", st.session_state.generated_plan, height=450)

st.header("4. Teacher Creative Adaptation")
teacher_adaptation = st.text_area(
    "Teacher modifications / creative adaptation",
    placeholder="Example: Make the challenge more suitable for 45 minutes, add a low-cost material option, simplify for weak students, connect to water/energy/sustainability in Qatar...",
    height=150,
)

if st.button("Apply Teacher Adaptation"):
    if not st.session_state.generated_plan:
        st.warning("Generate a plan first.")
    else:
        system_prompt = """
You are a STEM co-planning assistant. Revise the plan based on teacher creative adaptation.
Preserve curriculum grounding, STEM integration, realism, and teacher professional judgment.
"""
        user_prompt = f"""
ORIGINAL PLAN:
{st.session_state.generated_plan}

TEACHER CREATIVE ADAPTATION REQUEST:
{teacher_adaptation}

TASK:
Revise the plan accordingly. Clearly mark what changed based on teacher input.
"""
        with st.spinner("Applying teacher adaptation..."):
            st.session_state.generated_plan = call_ai(system_prompt, user_prompt, model=model)

st.header("5. Evaluate Plan Quality")
if st.button("Evaluate Plan Quality"):
    if not st.session_state.generated_plan:
        st.warning("Generate or paste a lesson plan first.")
    else:
        rubric = load_rubric()
        system_prompt = """
You are an AI-assisted evaluator of STEM lesson plan quality.
Your evaluation is preliminary and supports human evaluators. It does not replace expert judgment.
Use the rubric strictly. Provide transparent justifications for each score.
"""
        user_prompt = f"""
RUBRIC:
{rubric}

PLAN TO EVALUATE:
{st.session_state.generated_plan}

TASK:
Evaluate the plan using the rubric.
Provide:
1. Overall score out of 100
2. Sub-score table for each dimension out of 4
3. Evidence-grounding score
4. Completeness checklist
5. Strengths
6. Missing or weak components
7. Justifications for score deductions
8. Specific improvement suggestions
9. Final note reminding that this is AI-assisted preliminary evaluation
"""
        with st.spinner("Evaluating plan..."):
            st.session_state.evaluation = call_ai(system_prompt, user_prompt, model=model)

st.text_area("AI-Assisted Quality Evaluation", st.session_state.evaluation, height=350)

st.header("6. Export Report")
if st.button("Create Word Report"):
    metadata = {
        "Grade": grade,
        "Lesson title": lesson_title,
        "Duration": duration,
        "Class size": class_size,
        "Student level": student_level,
        "Resources": resources,
        "Local context": local_context,
        "Plan type": plan_type,
    }
    output_path = create_word_report(st.session_state.generated_plan, st.session_state.evaluation, metadata)
    with open(output_path, "rb") as f:
        st.download_button(
            "Download Word Report",
            data=f,
            file_name=output_path.name,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

st.divider()
st.caption("Research note: This prototype is designed to standardize AI-supported STEM planning, not to replace teacher professional judgment.")
